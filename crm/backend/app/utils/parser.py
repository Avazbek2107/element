"""
Word (.docx) va PDF fayllardan test savollarini parse qilish.

Qo'llab-quvvatlanadigan format:
    1. Savol matni?
    A) Variant A
    B) Variant B
    C) Variant C
    D) Variant D
    Javob: B

    yoki

    1-savol. Savol matni?
    A. Variant A
    B. Variant B
    C. Variant C
    D. Variant D
    To'g'ri javob: C
"""
import re
from typing import IO


def _parse_text(text: str) -> list[dict]:
    """Matndan savollar ro'yxatini ajratib oladi."""
    # Satrlarga ajratib, bo'sh satrlarni tozalaymiz
    lines = [ln.strip() for ln in text.splitlines()]
    lines = [ln for ln in lines if ln]

    questions = []
    current: dict | None = None

    # Savol boshi: "1.", "1)", "1-savol.", "Savol 1."
    q_start = re.compile(
        r'^(?:\d+[\.\-](?:savol)?\.?|\d+\))\s+(.+)', re.IGNORECASE
    )
    # Variantlar: "A)", "A.", "a)", "a."
    opt = re.compile(r'^([ABCDabcd])[\.)\-]\s+(.+)')
    # Javob satri: "Javob: B", "To'g'ri javob: A", "Answer: C"
    ans = re.compile(
        r"(?:to['']?g['']?ri\s+)?javob\s*[:=\-]\s*([ABCDabcd])",
        re.IGNORECASE
    )

    for line in lines:
        m = q_start.match(line)
        if m:
            if current and _is_complete(current):
                questions.append(current)
            current = {
                'question_text': m.group(1).strip(),
                'option_a': '',
                'option_b': '',
                'option_c': '',
                'option_d': '',
                'correct_answer': 'A',
                'points': 1,
            }
            continue

        if current is None:
            continue

        m = opt.match(line)
        if m:
            letter = m.group(1).upper()
            value = m.group(2).strip()
            key = f'option_{letter.lower()}'
            current[key] = value
            continue

        m = ans.search(line)
        if m:
            current['correct_answer'] = m.group(1).upper()
            continue

        # Savol matni ko'p satrda bo'lsa qo'shamiz
        if (
            not current.get('option_a')
            and not line[0:2].upper() in ('A)', 'A.', 'B)', 'B.')
        ):
            current['question_text'] += ' ' + line

    if current and _is_complete(current):
        questions.append(current)

    return questions


def _is_complete(q: dict) -> bool:
    return all([
        q.get('question_text'),
        q.get('option_a'),
        q.get('option_b'),
        q.get('option_c'),
        q.get('option_d'),
    ])


def parse_docx(file: IO[bytes]) -> list[dict]:
    from docx import Document
    doc = Document(file)
    text = '\n'.join(p.text for p in doc.paragraphs)
    return _parse_text(text)


def parse_pdf(file: IO[bytes]) -> list[dict]:
    import pdfplumber
    text_parts = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                text_parts.append(t)
    return _parse_text('\n'.join(text_parts))


XLSX_HEADERS = ["Savol", "A", "B", "C", "D", "Togri javob", "Ball"]


def parse_xlsx(file: IO[bytes]) -> list[dict]:
    """Excel fayldan savollarni o'qiydi. Ustunlar: Savol, A, B, C, D, Togri javob, Ball."""
    import openpyxl
    wb = openpyxl.load_workbook(file)
    ws = wb.active
    headers = [str(c.value).strip() if c.value else "" for c in ws[1]]
    label_map = {
        "Savol": "question_text", "A": "option_a", "B": "option_b",
        "C": "option_c", "D": "option_d",
        "Togri javob": "correct_answer", "To'g'ri javob": "correct_answer",
        "Ball": "points",
    }
    keys = [label_map.get(h, h.lower()) for h in headers]

    questions = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        if not any(row):
            continue
        item = {keys[i]: row[i] for i in range(len(keys)) if i < len(row)}
        q = {
            'question_text':  str(item.get('question_text') or '').strip(),
            'option_a':       str(item.get('option_a') or '').strip(),
            'option_b':       str(item.get('option_b') or '').strip(),
            'option_c':       str(item.get('option_c') or '').strip(),
            'option_d':       str(item.get('option_d') or '').strip(),
            'correct_answer': str(item.get('correct_answer') or 'A').strip().upper()[:1] or 'A',
            'points':         int(item.get('points') or 1),
        }
        if _is_complete(q):
            questions.append(q)
    return questions


def build_xlsx(questions: list) -> bytes:
    """Savollar ro'yxatidan Excel fayl (bytes) yaratadi. parse_xlsx bilan round-trip mos."""
    import openpyxl
    import io
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Savollar"
    ws.append(XLSX_HEADERS)
    for q in questions:
        ws.append([
            q.question_text, q.option_a, q.option_b, q.option_c, q.option_d,
            q.correct_answer, q.points,
        ])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def build_docx(test_title: str, questions: list) -> bytes:
    """Savollar ro'yxatidan Word fayl (bytes) yaratadi. parse_docx bilan round-trip mos."""
    from docx import Document
    import io
    doc = Document()
    doc.add_heading(test_title, level=1)
    for i, q in enumerate(questions, start=1):
        doc.add_paragraph(f"{i}. {q.question_text}")
        doc.add_paragraph(f"A) {q.option_a}")
        doc.add_paragraph(f"B) {q.option_b}")
        doc.add_paragraph(f"C) {q.option_c}")
        doc.add_paragraph(f"D) {q.option_d}")
        doc.add_paragraph(f"Javob: {q.correct_answer}")
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
